"""
ADGM Corporate Agent — Upgraded single-file prototype with offline RAG
Filename: ADGM_CorporateAgent_Main.py

Features:
- Gradio UI accepting .docx
- Better .docx parsing: paragraphs, tables, headers, footers
- Robust doc-type detection (content + filename + fuzzy matching)
- Checklist verification against embedded ADGM checklist
- Red-flag detection rules
- Offline RAG: builds TF-IDF index from local reference PDFs (page-level passages)
- Reviewer notes appended to .docx including RAG-cited passages and suggestions
- Outputs reviewed .docx and a structured JSON report
"""

import io
import os
import re
import json
import difflib
from typing import List, Dict, Tuple
from datetime import datetime

import gradio as gr
from docx import Document

# PDF reading & RAG (scikit-learn)
import PyPDF2
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# -----------------------------
# Configuration / References
# -----------------------------
REFERENCE_PDFS = ["Task.pdf", "Data Sources.pdf"]  # Make sure these files are present in working dir
RAG_TOP_K = 2  # top K passages to include as citations

ADGM_CHECKLISTS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors",
    ],
    # Extend as needed
}

DOC_KEYWORDS = {
    "articles of association": "Articles of Association",
    "memorandum of association": "Memorandum of Association",
    "memorandum of understanding": "Memorandum of Association",
    "moa": "Memorandum of Association",
    "aoa": "Articles of Association",
    "incorporation application": "Incorporation Application Form",
    "incorporation application form": "Incorporation Application Form",
    "ubo": "UBO Declaration Form",
    "ubo declaration": "UBO Declaration Form",
    "register of members": "Register of Members and Directors",
    "register of directors": "Register of Members and Directors",
}

# -----------------------------
# Utility: read PDFs into page-level passages
# -----------------------------
def build_rag_corpus(pdf_paths: List[str]) -> Tuple[List[str], List[Tuple[str,int]]]:
    """
    Returns (passages, metadata) where metadata[i] = (source_filename, page_index)
    Each passage corresponds to a PDF page's text (trimmed).
    """
    passages = []
    metadata = []
    for pdf_path in pdf_paths:
        if not os.path.exists(pdf_path):
            continue
        try:
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                for pidx in range(num_pages):
                    try:
                        page = reader.pages[pidx]
                        text = page.extract_text() or ""
                        text = text.strip()
                        if text:
                            # store the whole page as a passage (short documents)
                            passages.append(text)
                            metadata.append((os.path.basename(pdf_path), pidx + 1))  # 1-index page
                    except Exception:
                        continue
        except Exception:
            continue
    return passages, metadata

def build_tfidf_index(passages: List[str]):
    vec = TfidfVectorizer(stop_words="english", max_features=25000)
    if not passages:
        return vec, None
    X = vec.fit_transform(passages)
    return vec, X

def rag_query(query: str, vec: TfidfVectorizer, X, passages: List[str], metadata: List[Tuple[str,int]], top_k: int = 2):
    """
    Return top_k (passage_excerpt, metadata) for the query using TF-IDF cosine similarity.
    """
    if X is None or vec is None or not passages:
        return []
    qv = vec.transform([query])
    sims = cosine_similarity(qv, X).flatten()
    top_indices = sims.argsort()[::-1][:top_k]
    results = []
    for idx in top_indices:
        if sims[idx] <= 0:
            continue
        passage = passages[idx]
        src, page = metadata[idx]
        excerpt = passage.strip()
        # Keep excerpt reasonably short
        if len(excerpt) > 800:
            excerpt = excerpt[:800] + "..."
        results.append({"score": float(sims[idx]), "source": src, "page": page, "excerpt": excerpt})
    return results

# -----------------------------
# Better .docx parsing
# -----------------------------
def parse_docx_text_and_metadata(file_stream: io.BytesIO) -> Tuple[str, List[str]]:
    """
    Extract text from paragraphs, table cells, headers and footers.
    Return combined text and list of header/footer/table texts as metadata list.
    """
    file_stream.seek(0)
    doc = Document(file_stream)
    texts = []
    meta_texts = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            texts.append(p.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    meta_texts.append(cell.text.strip())
                    texts.append(cell.text.strip())

    # Headers & footers (python-docx header/footer exposure may vary by docx)
    try:
        for section in doc.sections:
            header = section.header
            footer = section.footer
            for p in header.paragraphs:
                if p.text and p.text.strip():
                    meta_texts.append(p.text.strip())
                    texts.append(p.text.strip())
            for p in footer.paragraphs:
                if p.text and p.text.strip():
                    meta_texts.append(p.text.strip())
                    texts.append(p.text.strip())
    except Exception:
        # Some documents may not expose headers/footers; ignore
        pass

    combined = "\n".join(texts)
    return combined, meta_texts

# -----------------------------
# Document detection (content + filename + fuzzy)
# -----------------------------
def detect_document_types(text: str, filename: str = "") -> List[str]:
    found = set()
    lowered = (text or "").lower()
    name_lower = (filename or "").lower()

    # Exact keyword matches in text
    for k, v in DOC_KEYWORDS.items():
        if k in lowered:
            found.add(v)

    # Filename clues
    for k, v in DOC_KEYWORDS.items():
        if k in name_lower:
            found.add(v)

    # Fuzzy: check each canonical doc name against text and filename using difflib
    canonical_names = list(set(DOC_KEYWORDS.values()))
    for cname in canonical_names:
        # build small phrases to search
        cname_lower = cname.lower()
        # if a close match in text or filename
        if difflib.get_close_matches(cname_lower, [lowered], n=1, cutoff=0.7):
            found.add(cname)
        if difflib.get_close_matches(cname_lower, [name_lower], n=1, cutoff=0.6):
            found.add(cname)

    return sorted(found)

def detect_process_from_uploaded(docs_detected: List[str]) -> str:
    # If any doc matches incorporation checklist, assume incorporation
    incorporation_set = set(ADGM_CHECKLISTS.get("Company Incorporation", []))
    if any(d in incorporation_set for d in docs_detected):
        return "Company Incorporation"
    return "Unknown"

# -----------------------------
# Checklist verification
# -----------------------------
def verify_checklist(process: str, docs_detected: List[str]) -> Dict:
    if process not in ADGM_CHECKLISTS:
        return {
            "process": process,
            "documents_uploaded": len(docs_detected),
            "required_documents": None,
            "missing_documents": [],
        }
    required = ADGM_CHECKLISTS[process]
    missing = [r for r in required if r not in docs_detected]
    return {
        "process": process,
        "documents_uploaded": len(docs_detected),
        "required_documents": len(required),
        "missing_documents": missing,
    }

# -----------------------------
# Red-flag detection (rules) - keep these simple but actionable
# -----------------------------
def detect_red_flags(text: str) -> List[Dict]:
    issues = []
    lowered = (text or "").lower()

    # Jurisdiction wrong
    if re.search(r"uae federal courts|federal courts of the uae|u\.a\.e\. federal courts", lowered):
        issues.append({
            "section": "Jurisdiction clause",
            "issue": "References UAE Federal Courts instead of ADGM",
            "severity": "High",
            "suggestion": "Update jurisdiction to ADGM Courts as per ADGM Companies Regulations."
        })

    # Missing signatory
    if not re.search(r"(signature\s*:|signed by|signature\b|signatory)", lowered):
        issues.append({
            "section": "Signatory",
            "issue": "No clear signatory or signature block found",
            "severity": "Medium",
            "suggestion": "Add signatory block with name, position, date, and signature."
        })

    # Ambiguous phrasing
    ambiguous_patterns = [
        (r"\bmay\b", "may"),
        (r"endeavour to", "endeavour to"),
        (r"best endeavor", "best endeavor"),
        (r"best endeavour", "best endeavour"),
        (r"best efforts", "best efforts"),
    ]
    for pattern, display in ambiguous_patterns:
        if re.search(pattern, lowered):
            issues.append({
                "section": "Wording",
                "issue": f"Ambiguous or discretionary language found: '{display}'",
                "severity": "Low",
                "suggestion": "Consider using 'shall' or 'must' for obligations."
            })
            break

    # UBO missing
    if "ubo" not in lowered and "ultimate beneficial" not in lowered:
        issues.append({
            "section": "UBO",
            "issue": "No UBO (Ultimate Beneficial Owner) declaration found",
            "severity": "High",
            "suggestion": "Include a UBO declaration in accordance with ADGM requirements."
        })

    # Non-ADGM template like DIFC
    if re.search(r"difc|dubai international financial centre", lowered):
        issues.append({
            "section": "Template/Jurisdiction",
            "issue": "Document references DIFC or other non-ADGM jurisdiction/template",
            "severity": "High",
            "suggestion": "Use ADGM-specific templates and guidance from ADGM portal."
        })

    return issues

# -----------------------------
# Annotate .docx: append reviewer notes including RAG citations
# -----------------------------
def annotate_docx_with_comments(original_stream: io.BytesIO, issues: List[Dict], rag_results_for_issue: Dict[int, List[Dict]]) -> io.BytesIO:
    """
    Appends a review section at the end of the document.
    For each issue, includes: issue details, suggestion, and RAG citations (source + page + excerpt).
    """
    original_stream.seek(0)
    doc = Document(original_stream)

    doc.add_page_break()
    doc.add_paragraph("--- ADGM Corporate Agent Review Notes ---")
    doc.add_paragraph(f"Generated on: {datetime.utcnow().isoformat()} UTC")

    if not issues:
        doc.add_paragraph("No issues detected by automated checks.")
    else:
        for idx, issue in enumerate(issues, start=1):
            doc.add_paragraph(f"Issue {idx}: {issue.get('section', 'Unknown Section')}")
            doc.add_paragraph(f"Issue detail: {issue.get('issue')}")
            doc.add_paragraph(f"Severity: {issue.get('severity')}")
            doc.add_paragraph(f"Suggested fix: {issue.get('suggestion')}")
            # RAG citations
            rag_entries = rag_results_for_issue.get(idx - 1, [])
            if rag_entries:
                doc.add_paragraph("Relevant ADGM references (top matches):")
                for e in rag_entries:
                    src = e.get("source")
                    page = e.get("page")
                    score = e.get("score")
                    excerpt = e.get("excerpt")
                    doc.add_paragraph(f"- Source: {src} (page {page}), score={score:.3f}")
                    # insert a short excerpt; keep it on a single paragraph
                    excerpt_para = excerpt.replace("\n", " ")
                    doc.add_paragraph(f"  Excerpt: {excerpt_para}")
            else:
                doc.add_paragraph("No relevant ADGM references found in local corpus.")
            doc.add_paragraph("")  # spacing

    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream

# -----------------------------
# Top-level processing pipeline
# -----------------------------
def process_uploaded_docx_bytes(file_bytes: bytes, filename: str,
                                vec, X, passages, metadata) -> Tuple[io.BytesIO, Dict]:
    # parse docx
    stream = io.BytesIO(file_bytes)
    combined_text, meta_texts = parse_docx_text_and_metadata(stream)

    # detect document types
    detected_docs = detect_document_types(combined_text, filename)

    # detect process
    process = detect_process_from_uploaded(detected_docs)

    # checklist
    checklist_result = verify_checklist(process, detected_docs)

    # red flags
    issues = detect_red_flags(combined_text)

    # For each detected issue, run RAG query to find supporting ADGM references
    rag_results_for_issue = {}
    for idx, issue in enumerate(issues):
        q = f"{issue.get('issue')} {issue.get('section')} {issue.get('suggestion')}"
        matches = rag_query(q, vec, X, passages, metadata, top_k=RAG_TOP_K)
        rag_results_for_issue[idx] = matches

    # annotate docx with comments + rag citations
    reviewed_stream = annotate_docx_with_comments(stream, issues, rag_results_for_issue)

    report = {
        "process": checklist_result.get("process", process),
        "documents_detected": detected_docs,
        "documents_uploaded": checklist_result.get("documents_uploaded"),
        "required_documents": checklist_result.get("required_documents"),
        "missing_documents": checklist_result.get("missing_documents"),
        "issues_found": issues,
        "rag_matches": {str(i): rag_results_for_issue.get(i, []) for i in range(len(issues))}
    }

    return reviewed_stream, report

# -----------------------------
# Build RAG index at startup
# -----------------------------
print("Building offline RAG index from local PDFs...")
passages, metadata = build_rag_corpus(REFERENCE_PDFS)
if passages:
    vec, X = build_tfidf_index(passages)
    print(f"Indexed {len(passages)} passages from local PDFs.")
else:
    vec, X = None, None
    print("No reference PDF passages found. RAG will be disabled.")

# -----------------------------
# Gradio UI
# -----------------------------
def gradio_interface():
    with gr.Blocks() as demo:
        gr.Markdown("# ADGM Corporate Agent — Upgraded Prototype\nUpload a .docx legal document to run checks, receive ADGM-cited reviewer notes and a JSON report.")
        doc_in = gr.File(label="Upload .docx file", file_types=[".docx"])
        output_doc = gr.File(label="Reviewed .docx (download)")
        json_out = gr.JSON(label="Structured Report")
        run_btn = gr.Button("Run Review")

        def on_run(uploaded):
            if not uploaded:
                return None, {"error": "No file uploaded"}
            # uploaded may be a tempfile-like object or dict depending on Gradio version
            try:
                path = uploaded.name
            except Exception:
                path = uploaded.get("name") if isinstance(uploaded, dict) else None
            if not path or not os.path.exists(path):
                # gradio may store file under 'tmp' path; attempt to read bytes directly
                try:
                    b = uploaded.read()
                    filename = getattr(uploaded, "name", "uploaded.docx")
                except Exception:
                    return None, {"error": "Cannot read uploaded file."}
            else:
                with open(path, "rb") as f:
                    b = f.read()
                filename = os.path.basename(path)

            reviewed_stream, report = process_uploaded_docx_bytes(b, filename, vec, X, passages, metadata)

            out_name = f"reviewed_{int(datetime.utcnow().timestamp())}_{filename}"
            with open(out_name, "wb") as wf:
                wf.write(reviewed_stream.read())

            return out_name, report

        run_btn.click(on_run, inputs=[doc_in], outputs=[output_doc, json_out])

    return demo

if __name__ == "__main__":
    app = gradio_interface()
    app.launch()
